"""Génère Consigne_Entrainement.docx — version entraînement du CT (lundi)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLEU_SORBONNE = RGBColor(0x1A, 0x27, 0x44)
OR_SORBONNE = RGBColor(0xC8, 0xA4, 0x15)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(20)
    run.font.color.rgb = BLEU_SORBONNE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = BLEU_SORBONNE

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(12)
    run.font.color.rgb = OR_SORBONNE

def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(11)
    if bold: r.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(11)

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0);  section.bottom_margin = Cm(2.0)

add_h1(doc, "Entraînement au Contrôle Terminal — SIASS")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("L3 Sciences Sociales — Séance d'entraînement du lundi")
r.italic = True; r.font.size = Pt(11)

# --- Préambule entraînement ---
add_h2(doc, "Objectif de la séance")
add_p(doc,
      "Cette séance d'entraînement est construite sur le même jeu de données et "
      "la même structure de questions que le Contrôle Terminal. Seules les "
      "variables croisées changent. L'objectif est de vous permettre de rejouer "
      "en conditions réelles la démarche attendue le jour de l'épreuve, avant "
      "d'aborder le CT officiel.")
add_p(doc,
      "Un corrigé complet de cette séance vous sera fourni : utilisez-le comme "
      "référence du format attendu (tableaux, graphiques, rédaction, "
      "interprétations).")

# --- Consignes générales ---
add_h2(doc, "Consignes")
add_p(doc,
      "Réalisez en deux heures l'ensemble des questions ci-dessous. Il est "
      "conseillé de passer 1 heure sur la Partie 1 (calculs) et 1 heure sur la "
      "Partie 2 (analyse rédigée).")
add_p(doc,
      "Toutes les ressources du cours sont accessibles. Aucune communication "
      "entre étudiants ; les seules communications permises sont avec "
      "l'enseignant.")

# --- Données ---
add_h2(doc, "Présentation des données")
add_p(doc,
      "Vous travaillez sur le jeu de données « MesHéritiers » (G. Henri-Panabière, "
      "1999) : 677 questionnaires recueillis auprès de collégiens lyonnais de "
      "6ème et 5ème issus majoritairement de familles fortement diplômées.")
add_bullet(doc, "Donnees_CT.xlsx — données brutes (1 ligne = 1 collégien).")
add_bullet(doc, "Dico_var_CT.docx — dictionnaire des variables.")

# --- Hypothèse ---
add_h2(doc, "Hypothèse à vérifier")
p = doc.add_paragraph()
r = p.add_run(
    "Au sein de familles fortement diplômées, le talent scolaire est-il "
    "transversal entre disciplines et le sentiment d'aisance avec les "
    "camarades de classe prédit-il les difficultés scolaires des "
    "collégiens ?"
)
r.italic = True; r.bold = True; r.font.size = Pt(11)
add_p(doc,
      "Cette hypothèse s'inscrit dans la lignée des travaux de G. Henri-Panabière "
      "sur les « héritiers en difficulté » : vous devez la mobiliser pour "
      "structurer votre analyse, en articulant les résultats de la Partie 1.")

# --- Partie 1 ---
add_h2(doc, "Partie 1 — Calculs et indicateurs (12 points)")
add_p(doc, "À l'aide de R ou d'Excel.", bold=True)

add_h3(doc, "Q1 — Statistiques descriptives (3 pts)")
add_p(doc,
      "Pour les variables V0082 (moyenne en biologie) et V0087 (moyenne en "
      "technologie), calculez les huit indicateurs suivants :")
add_bullet(doc, "minimum, premier quartile (Q1), médiane, troisième quartile (Q3), maximum,")
add_bullet(doc, "écart inter-quartile (EIQ), variance, écart-type.")
add_p(doc,
      "Présentez les résultats dans un tableau soigné et accompagnez chaque "
      "indicateur d'un court commentaire explicatif (ce qu'il mesure et ce que "
      "sa valeur révèle sur la distribution observée).")

add_h3(doc, "Q2 — Indicateurs bivariés (1 pt)")
add_p(doc,
      "Calculez la covariance, le coefficient de corrélation r et le coefficient "
      "de détermination R² entre la moyenne en biologie (V0082) et la moyenne en "
      "technologie (V0087). Expliquez brièvement ce que mesure chacun des trois "
      "indicateurs et ce qu'il nous dit ici.")

add_h3(doc, "Q3 — Régression linéaire (2 pts)")
add_p(doc,
      "Effectuez une régression linéaire de la moyenne en biologie (V0082) en "
      "fonction de la moyenne en technologie (V0087). Donnez l'équation "
      "Y = aX + b et expliquez ce que représente chacun des termes Y, a, X et b.")

add_h3(doc, "Q4 — Tableau croisé (2 pts)")
add_p(doc,
      "Construisez le tableau croisé entre la variable V0001 (sentiment de "
      "l'enfant face à ses camarades) et la variable V0008 (recodage des "
      "difficultés scolaires actuelles). Présentez deux tableaux soignés :")
add_bullet(doc, "le tableau des effectifs observés (avec totaux marginaux) ;")
add_bullet(doc, "le tableau des pourcentages en ligne.")
add_p(doc,
      "Proposez ensuite une courte lecture du tableau des pourcentages en ligne "
      "en commentant au moins une proportion qui vous paraît parlante.")

add_h3(doc, "Q5 — Test du Khi-deux (2 pts)")
add_p(doc,
      "À partir du tableau d'effectifs construit en Q4, calculez la p-value du "
      "test du Khi-deux (formule Excel : =TEST.CHISQ(...) ou, sous R, la fonction "
      "chisq.test()). Concluez : les deux variables sont-elles dépendantes ou "
      "indépendantes ? Justifiez en comparant la p-value obtenue au seuil "
      "standard de 0,05.")

add_h3(doc, "Q6 — Résidus standardisés et visualisation (2 pts)")
add_p(doc,
      "Calculez le tableau des résidus standardisés (formule : (Observé − Théorique) / √Théorique) "
      "et représentez-les sous forme de graphique en barres. Proposez une courte "
      "lecture de ce graphique.")

# --- Partie 2 ---
add_h2(doc, "Partie 2 — Analyse rédigée (6 points)")
add_p(doc,
      "Adoptez les standards d'un institut d'études. Chaque résultat statistique "
      "doit être accompagné d'une interprétation. Votre rédaction doit "
      "comporter au minimum :")
add_bullet(doc, "une introduction présentant le contexte de l'enquête ;")
add_bullet(doc, "le rappel de la problématique / hypothèse ;")
add_bullet(doc, "un développement mobilisant les Q1 à Q6 ;")
add_bullet(doc, "une conclusion qui répond à l'hypothèse et discute les limites.")

# --- Présentation ---
add_h2(doc, "Présentation et livrable (2 points)")
add_p(doc, "Formats privilégiés : PDF ou HTML. Format toléré : Word (.docx), mais privilégiez l'export PDF.")
add_p(doc, "Évaluation : lisibilité des tableaux (en-têtes, alignements, unités), légende des graphiques (titres, axes, échelles), rigueur typographique.")

# --- Bonus ---
add_h2(doc, "Questions bonus (2 points)")
add_p(doc,
      "Les deux questions ci-dessous sont optionnelles. Elles ne sont pas "
      "comptabilisées dans le barème principal mais peuvent être ajoutées à la "
      "note finale, qui peut ainsi atteindre 22/20.")

add_h3(doc, "B1 — V de Cramer (1 pt)")
add_p(doc,
      "À partir du Khi-deux calculé en Q5, calculez le V de Cramer (formule : "
      "V = √(Khi² / (n × min(k−1, l−1))), où k et l sont le nombre de modalités "
      "des deux variables). Interprétez sa valeur (< 0,10 = effet nul/très "
      "faible ; 0,10–0,20 = faible ; 0,20–0,30 = modéré ; > 0,30 = fort).")

add_h3(doc, "B2 — PEM (Pourcentage de l'Écart Maximum) (1 pt)")
add_p(doc,
      "Identifiez la case du tableau croisé Q4 qui présente la plus forte "
      "attraction (résidu positif maximal). Calculez son PEM selon la formule "
      "de Ph. Cibois :")
add_bullet(doc, "écart à l'indépendance = Observé − Théorique")
add_bullet(doc, "effectif maximum compatible avec les marges = min(Total Ligne, Total Colonne)")
add_bullet(doc, "écart maximum = effectif maximum − Théorique")
add_bullet(doc, "PEM = (écart à l'indépendance / écart maximum) × 100")
add_p(doc, "Interprétez la valeur obtenue (0 % = indépendance ; 100 % = attraction maximale).")

doc.save(r"C:\Users\oflah\SIASS-L3S6\CT\Consigne_Entrainement.docx")
print("OK Consigne_Entrainement.docx generated")
