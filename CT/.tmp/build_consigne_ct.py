"""Génère Consigne_CT.docx (version officielle du Contrôle Terminal)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLEU_SORBONNE = RGBColor(0x1A, 0x27, 0x44)
OR_SORBONNE = RGBColor(0xC8, 0xA4, 0x15)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BLEU_SORBONNE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLEU_SORBONNE

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = OR_SORBONNE

def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)

doc = Document()

# Marges
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

add_h1(doc, "Contrôle Terminal — SIASS")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("L3 Sciences Sociales — Université Paris 1 Panthéon-Sorbonne")
r.italic = True
r.font.size = Pt(11)

# === Consignes générales ===
add_h2(doc, "Consignes générales")
add_p(doc,
      "Réalisez en deux heures l'ensemble des questions ci-dessous. Vous pouvez "
      "organiser votre temps librement, mais il est vivement conseillé de passer "
      "1 heure sur la première partie (calculs) et 1 heure sur la seconde partie "
      "(analyse rédigée).")
add_p(doc,
      "Toutes les ressources utilisées en cours sont accessibles durant la durée "
      "de l'épreuve. Cependant, aucune communication entre étudiants n'est permise, "
      "de quelque manière que ce soit. Les seules communications permises sont avec "
      "l'enseignant.")

# === Données ===
add_h2(doc, "Présentation des données")
add_p(doc,
      "Vous travaillez sur le jeu de données « MesHéritiers » (G. Henri-Panabière, "
      "1999), constitué de 677 questionnaires recueillis auprès de collégiens "
      "scolarisés en 6ème et 5ème dans quatre établissements lyonnais accueillant "
      "majoritairement des familles fortement diplômées.")
add_p(doc, "Les fichiers à votre disposition sont :")
add_bullet(doc, "Donnees_CT.xlsx — données brutes (1 ligne = 1 collégien, 1 colonne = 1 variable).")
add_bullet(doc, "Dico_var_CT.docx — dictionnaire des variables, indispensable pour interpréter les codes.")

# === Hypothèse ===
add_h2(doc, "Hypothèse à vérifier")
p = doc.add_paragraph()
r = p.add_run(
    "Au sein de familles fortement diplômées, le talent scolaire est-il "
    "transversal entre disciplines et le sentiment d'aisance face aux "
    "enseignants prédit-il les difficultés scolaires des collégiens ?"
)
r.italic = True
r.bold = True
r.font.size = Pt(11)
add_p(doc,
      "C'est cette hypothèse, formulée à partir des travaux de G. Henri-Panabière "
      "sur les « héritiers en difficulté », qui doit guider votre analyse. Les "
      "questions de la Partie 1 vous fournissent les éléments statistiques ; la "
      "Partie 2 vous demande de les mobiliser pour répondre à cette hypothèse.")

# === Partie 1 ===
add_h2(doc, "Partie 1 — Calculs et indicateurs (12 points)")
add_p(doc, "À l'aide de R ou d'Excel.", bold=True)

add_h3(doc, "Q1 — Statistiques descriptives (3 pts)")
add_p(doc,
      "Pour les variables V0078 (moyenne en français) et V0084 (moyenne en "
      "mathématiques), calculez les huit indicateurs suivants :")
add_bullet(doc, "minimum, premier quartile (Q1), médiane, troisième quartile (Q3), maximum,")
add_bullet(doc, "écart inter-quartile (EIQ), variance, écart-type.")
add_p(doc,
      "Présentez les résultats dans un tableau soigné, et accompagnez chaque "
      "indicateur d'un court commentaire explicatif (ce qu'il mesure et ce que "
      "sa valeur révèle sur la distribution observée).")

add_h3(doc, "Q2 — Indicateurs bivariés (1 pt)")
add_p(doc,
      "Calculez la covariance, le coefficient de corrélation r et le coefficient "
      "de détermination R² entre la moyenne en français (V0078) et la moyenne "
      "en mathématiques (V0084). Expliquez brièvement ce que mesure chacun "
      "des trois indicateurs et ce qu'il nous dit ici.")

add_h3(doc, "Q3 — Régression linéaire (2 pts)")
add_p(doc,
      "Effectuez une régression linéaire de la moyenne en français (V0078) en "
      "fonction de la moyenne en mathématiques (V0084). Donnez l'équation "
      "Y = aX + b et expliquez ce que représente chacun des termes Y, a, X et b.")

add_h3(doc, "Q4 — Tableau croisé (2 pts)")
add_p(doc,
      "Construisez le tableau croisé entre la variable V0002 (sentiment de "
      "l'enfant face à ses enseignants) et la variable V0008 (recodage des "
      "difficultés scolaires actuelles). Présentez deux tableaux soignés :")
add_bullet(doc, "le tableau des effectifs observés (avec totaux marginaux) ;")
add_bullet(doc, "le tableau des pourcentages en ligne.")
add_p(doc,
      "Proposez ensuite une courte lecture du tableau des pourcentages en ligne "
      "en commentant au moins une proportion qui vous paraît parlante.")

add_h3(doc, "Q5 — Test du Khi-deux (2 pts)")
add_p(doc,
      "À partir du tableau d'effectifs construit en Q4, calculez la p-value du "
      "test du Khi-deux (formule Excel : =TEST.CHISQ(...) ou, sous R, la fonction "
      "chisq.test()). Concluez : les deux variables sont-elles dépendantes ou "
      "indépendantes ? Justifiez en comparant la p-value obtenue au seuil "
      "standard de 0,05.")

add_h3(doc, "Q6 — Résidus standardisés et visualisation (2 pts)")
add_p(doc,
      "Calculez le tableau des résidus standardisés (formule : (Observé − Théorique) / √Théorique) "
      "et représentez-les sous forme de graphique en barres. Proposez une courte "
      "lecture de ce graphique.")

# === Partie 2 ===
add_h2(doc, "Partie 2 — Analyse rédigée (6 points)")
add_p(doc,
      "Pour cette restitution, adoptez les standards d'un institut d'études. Ne "
      "vous contentez pas de juxtaposer des chiffres bruts ou des graphiques : "
      "interprétez vos résultats. Chaque sortie statistique ou cartographique "
      "doit être accompagnée d'une analyse claire mettant en lumière les "
      "dynamiques mises en évidence dans la Partie 1.")
add_p(doc, "Votre rédaction doit comporter au minimum :")
add_bullet(doc, "une introduction présentant le contexte de l'enquête (G. Henri-Panabière, 1999) ;")
add_bullet(doc, "le rappel explicite de la problématique / hypothèse ;")
add_bullet(doc, "un développement structuré qui mobilise les résultats des Q1 à Q6 ;")
add_bullet(doc, "une conclusion qui répond à l'hypothèse et discute les limites de l'analyse.")

# === Présentation ===
add_h2(doc, "Présentation et livrable (2 points)")
add_p(doc, "Formats privilégiés : PDF ou HTML (recommandés pour figer la mise en page, "
            "particulièrement si vous générez votre rapport via R/Quarto).")
add_p(doc, "Format toléré : Word (.docx) accepté, mais à vos risques et périls (la mise "
            "en forme peut être altérée). Privilégiez l'export PDF.")
add_p(doc, "Vous serez évalué·e sur la lisibilité de vos tableaux (en-têtes, alignements, "
            "unités), la légende de vos graphiques (titres, axes, échelles) et la "
            "rigueur typographique générale du document.")

# === Bonus ===
add_h2(doc, "Questions bonus (2 points)")
add_p(doc,
      "Les deux questions ci-dessous sont optionnelles. Elles ne sont pas "
      "comptabilisées dans le barème principal mais peuvent être ajoutées à la "
      "note finale, qui peut ainsi atteindre 22/20.")

add_h3(doc, "B1 — V de Cramer (1 pt)")
add_p(doc,
      "À partir du Khi-deux calculé en Q5, calculez le V de Cramer (formule : "
      "V = √(Khi² / (n × min(k−1, l−1))), où k et l sont le nombre de modalités "
      "des deux variables). Interprétez sa valeur en fonction des seuils utilisés "
      "en sciences sociales (effet nul ou très faible si V < 0,10 ; faible "
      "0,10–0,20 ; modéré 0,20–0,30 ; fort > 0,30).")

add_h3(doc, "B2 — PEM (Pourcentage de l'Écart Maximum) (1 pt)")
add_p(doc,
      "Identifiez la case du tableau croisé Q4 qui présente la plus forte "
      "attraction (résidu positif maximal). Calculez son PEM selon la formule "
      "de Ph. Cibois :")
add_bullet(doc, "écart à l'indépendance = Observé − Théorique")
add_bullet(doc, "effectif maximum compatible avec les marges = min(Total Ligne, Total Colonne)")
add_bullet(doc, "écart maximum = effectif maximum − Théorique")
add_bullet(doc, "PEM = (écart à l'indépendance / écart maximum) × 100")
add_p(doc,
      "Interprétez la valeur obtenue (0 % = indépendance ; 100 % = attraction "
      "maximale possible compte tenu des marges).")

# === Modalités de dépôt ===
add_h2(doc, "Modalités de dépôt")
add_p(doc, "Vous devez impérativement déposer deux éléments sur l'EPI (Espace Pédagogique Interactif) :")
add_bullet(doc, "votre document de restitution final (PDF ou HTML) ;")
add_bullet(doc, "votre fichier de travail brut (fichier Excel ou script R).")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("L'heure limite de dépôt est fixée à 13h00 au plus tard.")
r.bold = True
r.font.size = Pt(12)

doc.save(r"C:\Users\oflah\SIASS-L3S6\CT\Consigne_CT.docx")
print("OK Consigne_CT.docx generated")
